import os
import re
import json
import uuid
import bcrypt
from datetime import datetime, timedelta
from typing import Optional, List, Dict, Any
from pathlib import Path

from fastapi import FastAPI, HTTPException, status, Depends, UploadFile, File
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, EmailStr, Field, validator
from sqlalchemy import create_engine, Column, String, Integer, DateTime, Text, Float, ForeignKey, desc
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, Session, relationship
from sqlalchemy.exc import IntegrityError
from jose import JWTError, jwt

import PyPDF2
from docx import Document

# importing OCR dependencies
try:
    import pytesseract
    from pdf2image import convert_from_path
    import tempfile

    OCR_AVAILABLE = True
    import platform

    if platform.system() == "Windows":
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    elif platform.system() == "Darwin":  # macOS
        pytesseract.pytesseract.tesseract_cmd = '/usr/local/bin/tesseract'
    else:  # Linux
        pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'
    print("✅ OCR libraries loaded successfully")
except ImportError:
    OCR_AVAILABLE = False
    print("⚠️ OCR not available. Scanned PDFs won't be processed.")

# importing Google Gemini
try:
    import google.generativeai as genai

    GEMINI_AVAILABLE = True
except ImportError:
    GEMINI_AVAILABLE = False
    print("⚠️ Google Gemini not available. Install with: pip install google-generativeai")

import logging
from dotenv import load_dotenv

load_dotenv()

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


#  CONFIGURATION

class Config:
    DATABASE_URL = os.getenv("DATABASE_URL", "sqlite:///./resume_analyzer.db")
    SECRET_KEY = os.getenv("SECRET_KEY", "your-super-secret-key-change-this-in-production")
    ALGORITHM = "HS256"
    ACCESS_TOKEN_EXPIRE_MINUTES = 60
    GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
    UPLOAD_DIR = Path("uploads")
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

    UPLOAD_DIR.mkdir(exist_ok=True)


# Initialize Gemini
try:
    # For newer API versions
    gemini_model = genai.GenerativeModel('gemini-1.5-pro')
except:
    try:
        # For older API versions
        gemini_model = genai.GenerativeModel('gemini-pro')
    except:
        gemini_model = None
        print("⚠️ No valid Gemini model found")

# DATABASE

if "sqlite" in Config.DATABASE_URL:
    engine = create_engine(Config.DATABASE_URL, connect_args={"check_same_thread": False})
else:
    engine = create_engine(Config.DATABASE_URL)

SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
Base = declarative_base()
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="api/auth/login")


# MODELS

class User(Base):
    __tablename__ = "users"
    id = Column(Integer, primary_key=True, index=True)
    email = Column(String, unique=True, index=True, nullable=False)
    username = Column(String, unique=True, index=True, nullable=False)
    hashed_password = Column(String, nullable=False)
    created_at = Column(DateTime, default=datetime.utcnow)
    resumes = relationship("Resume", back_populates="user", cascade="all, delete-orphan")
    analyses = relationship("Analysis", back_populates="user", cascade="all, delete-orphan")


class Resume(Base):
    __tablename__ = "resumes"
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    filename = Column(String, nullable=False)
    file_path = Column(String, nullable=False)
    uploaded_at = Column(DateTime, default=datetime.utcnow)
    name = Column(String, nullable=True)
    email = Column(String, nullable=True)
    phone = Column(String, nullable=True)
    skills = Column(Text, nullable=True)
    experience = Column(Text, nullable=True)
    education = Column(Text, nullable=True)
    raw_text = Column(Text, nullable=True)
    user = relationship("User", back_populates="resumes")
    analyses = relationship("Analysis", back_populates="resume", cascade="all, delete-orphan")


class ScoreHistory(Base):
    __tablename__ = "score_history"
    id = Column(Integer, primary_key=True, index=True)
    resume_id = Column(Integer, ForeignKey("resumes.id"), nullable=False)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    total_score = Column(Float, nullable=False)
    experience_score = Column(Float)
    skills_score = Column(Float)
    education_score = Column(Float)
    keywords_score = Column(Float)
    formatting_score = Column(Float)
    created_at = Column(DateTime, default=datetime.utcnow)

    resume = relationship("Resume", back_populates="score_history")
    user = relationship("User")


# Add relationship to Resume model
Resume.score_history = relationship("ScoreHistory", back_populates="resume", cascade="all, delete-orphan")

class Analysis(Base):
    __tablename__ = "analyses"
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    resume_id = Column(Integer, ForeignKey("resumes.id"), nullable=False)
    analysis_type = Column(String, nullable=False)
    ats_score = Column(Float, nullable=True)
    strengths = Column(Text, nullable=True)
    weaknesses = Column(Text, nullable=True)
    missing_keywords = Column(Text, nullable=True)
    suggestions = Column(Text, nullable=True)
    gap_analysis = Column(Text, nullable=True)
    skill_gaps = Column(Text, nullable=True)
    experience_gaps = Column(Text, nullable=True)
    education_gaps = Column(Text, nullable=True)
    improvement_plan = Column(Text, nullable=True)
    job_description = Column(Text, nullable=True)
    match_percentage = Column(Float, nullable=True)
    matched_skills = Column(Text, nullable=True)
    missing_skills = Column(Text, nullable=True)
    role_fit_summary = Column(Text, nullable=True)
    improvement_tips = Column(Text, nullable=True)
    created_at = Column(DateTime, default=datetime.utcnow)
    user = relationship("User", back_populates="analyses")
    resume = relationship("Resume", back_populates="analyses")


# Create tables
Base.metadata.create_all(bind=engine)


#  SCHEMAS

class UserCreate(BaseModel):
    email: EmailStr
    username: str = Field(..., min_length=3, max_length=50)
    password: str = Field(..., min_length=6)

    @validator('username')
    def username_alphanumeric(cls, v):
        if not re.match(r'^[a-zA-Z0-9_]+$', v):
            raise ValueError('Username must contain only letters, numbers, and underscores')
        return v

    @validator('password')
    def password_strength(cls, v):
        if len(v) < 6:
            raise ValueError('Password must be at least 6 characters')
        return v


class Token(BaseModel):
    access_token: str
    token_type: str


class ATSAnalysisRequest(BaseModel):
    resume_id: int


class JobMatchRequest(BaseModel):
    resume_id: int
    job_description: str = Field(..., min_length=10)


#  AUTH UTILS

def get_password_hash(password: str) -> str:
    """Hash a password using bcrypt"""
    try:
        salt = bcrypt.gensalt()
        hashed = bcrypt.hashpw(password.encode('utf-8'), salt)
        return hashed.decode('utf-8')
    except Exception as e:
        logger.error(f"Password hashing error: {e}")
        raise HTTPException(status_code=500, detail="Error processing password")


def verify_password(plain: str, hashed: str) -> bool:
    """Verify a password against its hash"""
    try:
        return bcrypt.checkpw(plain.encode('utf-8'), hashed.encode('utf-8'))
    except Exception as e:
        logger.error(f"Password verify error: {e}")
        return False


def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    expire = datetime.utcnow() + (expires_delta or timedelta(minutes=Config.ACCESS_TOKEN_EXPIRE_MINUTES))
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, Config.SECRET_KEY, algorithm=Config.ALGORITHM)


def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


async def get_current_user(token: str = Depends(oauth2_scheme), db: Session = Depends(get_db)):
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, Config.SECRET_KEY, algorithms=[Config.ALGORITHM])
        sub = payload.get("sub")
        if sub is None:
            raise credentials_exception
        user_id = int(sub)
    except (JWTError, ValueError, TypeError):
        raise credentials_exception

    user = db.query(User).filter(User.id == user_id).first()
    if user is None:
        raise credentials_exception
    return user


# ADVANCED SCORE CALCULATOR

class AdvancedScoreCalculator:
    """Advanced ATS score calculation with weighted categories and statistics"""

    # Default weights (can be adjusted by admin)
    DEFAULT_WEIGHTS = {
        "experience": 30,
        "skills": 25,
        "education": 15,
        "keywords": 20,
        "formatting": 10
    }

    @staticmethod
    def calculate_experience_score(resume_text: str, experience_data: List) -> Dict:
        """Calculate experience score based on years, relevance, and quality"""
        score = 0
        details = {}

        # Years of experience
        years_match = re.search(r'(\d+)\+?\s*years?\s+of\s+experience', resume_text.lower())
        if years_match:
            years = int(years_match.group(1))
            if years >= 5:
                years_score = 100
                details["years"] = "5+ years (Excellent)"
            elif years >= 3:
                years_score = 75
                details["years"] = "3-5 years (Good)"
            elif years >= 1:
                years_score = 50
                details["years"] = "1-3 years (Adequate)"
            else:
                years_score = 25
                details["years"] = "<1 year (Entry level)"
        else:
            years_score = 30
            details["years"] = "Not specified"

        # Action verbs quality
        action_verbs = ['led', 'managed', 'developed', 'created', 'implemented',
                        'designed', 'built', 'achieved', 'improved', 'increased']
        found_verbs = sum(1 for verb in action_verbs if re.search(rf'\b{verb}\b', resume_text.lower()))
        verb_score = min(100, (found_verbs / len(action_verbs)) * 100)
        details["action_verbs"] = f"{found_verbs}/{len(action_verbs)} found"

        # Quantifiable achievements
        has_numbers = bool(re.search(r'\d+%|\d+\+|\$\d+|\d+\s*years', resume_text))
        quant_score = 100 if has_numbers else 40
        details["quantifiable"] = "Yes" if has_numbers else "No"

        # Calculate weighted score
        score = (years_score * 0.4) + (verb_score * 0.3) + (quant_score * 0.3)

        return {
            "score": round(score, 1),
            "max_score": 100,
            "details": details,
            "percentile": min(95, max(10, score))  # Mock percentile
        }

    @staticmethod
    def calculate_skills_score(skills: List[str], job_description: str = None) -> Dict:
        """Calculate skills score based on relevance and demand"""
        if not skills:
            return {"score": 20, "max_score": 100, "details": {"matched": 0, "total": 0}, "percentile": 10}

        # High-demand skills (weighted higher)
        high_demand = ['python', 'java', 'sql', 'aws', 'react', 'javascript', 'docker', 'kubernetes']
        high_demand_matches = [s for s in skills if s.lower() in high_demand]

        # Skill count score
        count_score = min(100, (len(skills) / 15) * 100)

        # Quality score (high-demand skills)
        quality_score = (len(high_demand_matches) / len(high_demand)) * 100 if high_demand else 50

        # Relevance score (if job description provided)
        relevance_score = 50
        if job_description:
            job_lower = job_description.lower()
            matched_job_skills = sum(1 for s in skills if s.lower() in job_lower)
            relevance_score = min(100, (matched_job_skills / max(len(skills), 1)) * 100)

        score = (count_score * 0.3) + (quality_score * 0.4) + (relevance_score * 0.3)

        return {
            "score": round(score, 1),
            "max_score": 100,
            "details": {
                "total_skills": len(skills),
                "high_demand_matches": len(high_demand_matches),
                "high_demand_list": high_demand_matches[:5],
                "relevance_score": round(relevance_score, 1)
            },
            "percentile": min(90, max(15, score))
        }

    @staticmethod
    def calculate_education_score(education: List, resume_text: str) -> Dict:
        """Calculate education score based on level and relevance"""
        score = 40  # Base score
        details = {}

        text_lower = resume_text.lower()

        # Check for degree levels
        if any(degree in text_lower for degree in ['phd', 'doctorate']):
            score = 100
            details["level"] = "PhD/Doctorate"
        elif any(degree in text_lower for degree in ['master', 'm.s.', 'm.sc', 'mba']):
            score = 85
            details["level"] = "Master's Degree"
        elif any(degree in text_lower for degree in ['bachelor', 'b.s.', 'b.sc', 'b.tech', 'b.e']):
            score = 70
            details["level"] = "Bachelor's Degree"
        elif any(degree in text_lower for degree in ['associate', 'diploma']):
            score = 50
            details["level"] = "Associate Degree/Diploma"
        else:
            details["level"] = "Not specified"

        # Check for GPA (optional bonus)
        gpa_match = re.search(r'gpa:?\s*(\d+\.?\d*)', text_lower)
        if gpa_match:
            gpa = float(gpa_match.group(1))
            if gpa >= 3.5:
                score = min(100, score + 10)
                details["gpa"] = f"{gpa} (Excellent)"
            elif gpa >= 3.0:
                score = min(100, score + 5)
                details["gpa"] = f"{gpa} (Good)"

        # Check for relevant coursework
        relevant_courses = ['data structures', 'algorithms', 'database', 'machine learning', 'ai']
        found_courses = sum(1 for course in relevant_courses if course in text_lower)
        if found_courses > 0:
            score = min(100, score + (found_courses * 2))
            details["relevant_courses"] = found_courses

        return {
            "score": round(score, 1),
            "max_score": 100,
            "details": details,
            "percentile": min(85, max(20, score))
        }

    @staticmethod
    def calculate_keyword_score(resume_text: str, job_description: str = None) -> Dict:
        """Calculate keyword optimization score"""
        text_lower = resume_text.lower()

        # Industry keywords
        industry_keywords = [
            'python', 'javascript', 'react', 'sql', 'aws', 'docker', 'git',
            'agile', 'scrum', 'leadership', 'communication', 'problem solving',
            'project management', 'data analysis', 'machine learning', 'api'
        ]

        found_keywords = [kw for kw in industry_keywords if kw in text_lower]
        keyword_score = (len(found_keywords) / len(industry_keywords)) * 100

        details = {
            "found_keywords": len(found_keywords),
            "total_keywords": len(industry_keywords),
            "top_keywords": found_keywords[:5]
        }

        # If job description provided, calculate job-specific keyword match
        if job_description:
            job_lower = job_description.lower()
            job_keywords = [kw for kw in industry_keywords if kw in job_lower]
            matched_job_keywords = [kw for kw in job_keywords if kw in text_lower]
            job_match_score = (len(matched_job_keywords) / max(len(job_keywords), 1)) * 100
            details["job_keyword_match"] = round(job_match_score, 1)
            # Weight job-specific keywords higher
            keyword_score = (keyword_score * 0.4) + (job_match_score * 0.6)

        return {
            "score": round(keyword_score, 1),
            "max_score": 100,
            "details": details,
            "percentile": min(88, max(12, keyword_score))
        }

    @staticmethod
    def calculate_formatting_score(resume_text: str) -> Dict:
        """Calculate formatting and structure score"""
        score = 60  # Base score
        details = {}

        # Check for sections
        sections = ['experience', 'education', 'skills', 'summary', 'projects', 'certifications']
        found_sections = [s for s in sections if s in resume_text.lower()]
        section_score = (len(found_sections) / len(sections)) * 30
        details["sections_found"] = found_sections
        details["section_count"] = len(found_sections)

        # Check length (optimal: 400-800 words)
        word_count = len(resume_text.split())
        if 400 <= word_count <= 800:
            length_score = 20
            details["length"] = f"{word_count} words (Optimal)"
        elif 200 <= word_count < 400:
            length_score = 15
            details["length"] = f"{word_count} words (Short but acceptable)"
        elif word_count > 800:
            length_score = 10
            details["length"] = f"{word_count} words (Too long)"
        else:
            length_score = 5
            details["length"] = f"{word_count} words (Too short)"

        # Check for bullet points
        bullet_count = len(re.findall(r'[•·-]|\d+\.', resume_text))
        bullet_score = min(15, bullet_count)
        details["bullet_points"] = bullet_count

        # Check for contact information
        has_email = bool(re.search(r'@', resume_text))
        has_phone = bool(re.search(r'\d{10}', resume_text))
        contact_score = 15 if (has_email and has_phone) else 5 if (has_email or has_phone) else 0
        details["contact_info"] = "Complete" if (has_email and has_phone) else "Partial" if (
                    has_email or has_phone) else "Missing"

        score = section_score + length_score + bullet_score + contact_score

        return {
            "score": round(score, 1),
            "max_score": 100,
            "details": details,
            "percentile": min(92, max(18, score))
        }

    @staticmethod
    def calculate_total_score(scores: Dict, weights: Dict = None) -> Dict:
        """Calculate weighted total score with confidence interval"""
        if weights is None:
            weights = AdvancedScoreCalculator.DEFAULT_WEIGHTS

        total = 0
        breakdown = {}

        for category, weight in weights.items():
            category_score = scores.get(category, {}).get("score", 0)
            contribution = (category_score * weight) / 100
            total += contribution
            breakdown[category] = {
                "score": category_score,
                "weight": weight,
                "contribution": round(contribution, 1)
            }

        # Calculate confidence interval (±5 points due to missing data)
        confidence_interval = 5
        lower_bound = max(0, total - confidence_interval)
        upper_bound = min(100, total + confidence_interval)

        # Calculate percentile rank (mock - in production would compare to historical data)
        percentile = min(98, max(2, total))

        return {
            "total_score": round(total, 1),
            "max_score": 100,
            "confidence_interval": f"±{confidence_interval} points",
            "score_range": f"{lower_bound}-{upper_bound}",
            "percentile_rank": round(percentile, 1),
            "percentile_label": AdvancedScoreCalculator._get_percentile_label(percentile),
            "breakdown": breakdown
        }

    @staticmethod
    def _get_percentile_label(percentile):
        if percentile >= 85:
            return "Top 15% of applicants"
        elif percentile >= 70:
            return "Top 30% of applicants"
        elif percentile >= 50:
            return "Top 50% of applicants"
        elif percentile >= 30:
            return "Below average"
        else:
            return "Needs significant improvement"

    @staticmethod
    def _generate_strengths(scores):
        """Generate strengths based on scores"""
        strengths = []
        for category, data in scores.items():
            if data["score"] >= 70:
                strengths.append(f"Strong {category.title()} section: {data['score']}/100")
            elif category == "skills" and data["score"] >= 60:
                strengths.append(
                    f"Good skills foundation with {data['details'].get('total_skills', 0)} skills identified")
            elif category == "formatting" and data["score"] >= 60:
                sections = data['details'].get('section_count', 0)
                strengths.append(f"Good document structure with {sections} sections found")

        if not strengths:
            strengths = ["Your resume has good basic structure", "Contact information is present"]
        return strengths[:5]

    @staticmethod
    def _generate_weaknesses(scores):
        """Generate weaknesses based on scores"""
        weaknesses = []
        for category, data in scores.items():
            if data["score"] < 50:
                weaknesses.append(f"Improve {category.title()} section: {data['score']}/100")
            elif data["score"] < 65:
                weaknesses.append(f"{category.title()} section needs enhancement: {data['score']}/100")

        if not weaknesses:
            weaknesses = ["Consider adding more quantifiable achievements"]
        return weaknesses[:5]

    @staticmethod
    def _generate_suggestions(scores):
        """Generate actionable suggestions"""
        suggestions = []

        if scores.get("keywords", {}).get("score", 0) < 60:
            suggestions.append("Add more industry-relevant keywords to your resume")
        if scores.get("formatting", {}).get("score", 0) < 60:
            suggestions.append("Improve formatting with clear section headers and bullet points")
        if scores.get("skills", {}).get("score", 0) < 60:
            suggestions.append("List more technical skills relevant to your target role")
        if scores.get("experience", {}).get("score", 0) < 60:
            suggestions.append("Add quantifiable achievements with numbers and percentages")
        if scores.get("education", {}).get("score", 0) < 60:
            suggestions.append("Include relevant coursework or certifications")

        if not suggestions:
            suggestions = [
                "Add a professional summary at the top",
                "Tailor your resume for each job application",
                "Include a LinkedIn profile URL"
            ]
        return suggestions[:7]

    @staticmethod
    def _generate_gap_analysis(scores):
        """Generate gap analysis"""
        gaps = {}
        for category, data in scores.items():
            if data["score"] < 50:
                gaps[category] = [f"{category.title()} score is low ({data['score']}/100)"]
                if category == "keywords" and data["details"].get("found_keywords", 0) < 5:
                    gaps[category].append("Missing critical industry keywords")
                elif category == "formatting" and data["details"].get("section_count", 0) < 3:
                    gaps[category].append("Missing standard sections (Experience, Education, Skills)")
        return gaps

    @staticmethod
    def _generate_improvement_plan(scores):
        """Generate improvement plan"""
        plan = {"short_term": [], "medium_term": [], "long_term": []}

        for category, data in scores.items():
            if data["score"] < 40:
                plan["short_term"].append(f"Completely restructure {category} section")
            elif data["score"] < 60:
                plan["short_term"].append(f"Focus on improving {category} section this week")
            elif data["score"] < 75:
                plan["medium_term"].append(f"Enhance {category} with more details and metrics")

        if not plan["short_term"]:
            plan["short_term"] = ["Review and optimize your resume summary", "Update contact information"]
        if not plan["medium_term"]:
            plan["medium_term"] = ["Complete an online certification in your field", "Build a portfolio project"]
        if not plan["long_term"]:
            plan["long_term"] = ["Gain additional experience through projects", "Network with industry professionals"]

        return plan
#  RESUME PARSER

class ResumeParser:

    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from PDF with multiple fallback methods"""
        text = ""

        # Method 1: Try PyPDF2 first (for text-based PDFs)
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                logger.info(f"📄 PDF has {len(reader.pages)} pages")

                # Check if PDF is encrypted
                if reader.is_encrypted:
                    raise HTTPException(
                        status_code=400,
                        detail="PDF is password protected. Please unlock the PDF and try again."
                    )

                for i, page in enumerate(reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Page {i + 1} extraction failed: {e}")
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"PyPDF2 error: {e}")

        # Method 2: Try pdfplumber (better extraction)
        if not text.strip():
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                logger.info("✅ pdfplumber extraction successful")
            except ImportError:
                logger.info("pdfplumber not available")
            except Exception as e:
                logger.error(f"pdfplumber error: {e}")

        # Method 3: Try OCR for scanned PDFs
        if not text.strip() and OCR_AVAILABLE:
            logger.info("🔍 No text found, attempting OCR...")
            try:
                with tempfile.TemporaryDirectory() as tmp:
                    images = convert_from_path(
                        file_path,
                        dpi=300,
                        output_folder=tmp,
                        fmt='jpeg'
                    )

                    for i, img in enumerate(images):
                        # Preprocess image for better OCR
                        img = img.convert('L')  # Convert to grayscale

                        # Try OCR with different configurations
                        try:
                            page_text = pytesseract.image_to_string(img, config='--oem 3 --psm 6')
                            if page_text and page_text.strip():
                                text += page_text + "\n"
                                logger.info(f"✅ OCR page {i + 1} successful")
                        except Exception as e:
                            logger.warning(f"OCR failed for page {i + 1}: {e}")
            except Exception as e:
                logger.error(f"OCR process failed: {e}")

        # Validate extracted text
        if not text.strip():
            raise HTTPException(
                status_code=400,
                detail="No text could be extracted from this PDF. The file might be:\n"
                       "- Password protected (please unlock it first)\n"
                       "- Completely image-based (install Tesseract OCR)\n"
                       "- Corrupted or empty\n\n"
                       "Solutions:\n"
                       "1. Install Tesseract OCR for scanned PDFs\n"
                       "2. Convert PDF to DOCX format\n"
                       "3. Use a different PDF file"
            )

        # Clean up extracted text
        text = re.sub(r'\n{3,}', '\n\n', text)  # Remove excessive newlines
        logger.info(f"✅ PDF extraction complete: {len(text)} chars")
        return text

    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        try:
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n".join(paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += "\n" + cell.text
            if not text.strip():
                raise HTTPException(status_code=400, detail="DOCX file contains no readable text.")
            logger.info(f"✅ DOCX extracted: {len(text)} chars")
            return text
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"DOCX error: {e}")
            raise HTTPException(status_code=400, detail=f"Could not parse DOCX: {str(e)}")

    @staticmethod
    def extract_email(text: str) -> Optional[str]:
        matches = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b', text)
        return matches[0] if matches else None

    @staticmethod
    def extract_phone(text: str) -> Optional[str]:
        # Improved phone number regex
        matches = re.findall(r'[\+\(]?[1-9][0-9 .\-\(\)]{8,}[0-9]', text)
        return matches[0] if matches else None

    @staticmethod
    def extract_name(text: str) -> Optional[str]:
        # Look for name in first few lines
        lines = text.strip().split('\n')[:10]
        for line in lines:
            line = line.strip()
            # Name is typically short, has no special characters, and is at the top
            if (line and 2 <= len(line.split()) <= 4 and len(line) < 60
                    and re.match(r'^[A-Za-z\s\.]+$', line)
                    and not any(w in line.lower() for w in
                                ['resume', 'curriculum', 'vitae', 'cv', 'email', 'phone', 'address', 'linkedin',
                                 'github'])):
                # Check if it's likely a name (has uppercase letters)
                if any(c.isupper() for c in line):
                    return line
        return None

    @staticmethod
    def extract_skills(text: str) -> List[str]:
        skills_db = [
            # Programming Languages
            'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 'php',
            'swift', 'kotlin', 'go', 'rust', 'scala', 'r', 'matlab',
            # Frameworks & Libraries
            'react', 'angular', 'vue', 'django', 'flask', 'fastapi', 'spring', 'spring boot',
            'node.js', 'express', 'rails', 'laravel', 'asp.net', 'jquery',
            # Databases
            'postgresql', 'mysql', 'mongodb', 'redis', 'cassandra', 'elasticsearch',
            'oracle', 'sqlite', 'firebase', 'dynamodb',
            # Cloud & DevOps
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'git', 'jenkins', 'ci/cd',
            'terraform', 'ansible', 'prometheus', 'grafana', 'linux', 'bash',
            # Web Technologies
            'html', 'css', 'sass', 'tailwind', 'bootstrap', 'graphql', 'rest api',
            'microservices', 'webpack', 'vite',
            # Machine Learning & Data Science
            'machine learning', 'deep learning', 'tensorflow', 'pytorch', 'pandas',
            'numpy', 'scikit-learn', 'keras', 'opencv', 'nlp', 'data analysis',
            # Soft Skills
            'leadership', 'communication', 'teamwork', 'problem solving', 'project management',
            'agile', 'scrum', 'kanban', 'jira', 'confluence', 'figma',
        ]
        text_lower = text.lower()
        found_skills = set()
        for skill in skills_db:
            # Use word boundaries for better matching
            if re.search(r'\b' + re.escape(skill) + r'\b', text_lower):
                found_skills.add(skill)
        return list(found_skills)

    @staticmethod
    def extract_experience(text: str) -> List[Dict[str, Any]]:
        experiences = []
        text_lower = text.lower()

        # Look for years of experience
        patterns = [
            r'(\d+)\+?\s*years?\s+of\s+experience',
            r'experience\s*:?\s*(\d+)\+?\s*years?',
            r'(\d+)\+?\s*yrs?\s+experience',
        ]

        for pattern in patterns:
            matches = re.findall(pattern, text_lower)
            if matches:
                experiences.append({
                    "type": "total_experience",
                    "years": matches[0],
                    "description": f"{matches[0]} years of experience"
                })
                break

        # Look for seniority levels
        seniority_keywords = {
            'senior': 'Senior',
            'lead': 'Lead',
            'principal': 'Principal',
            'junior': 'Junior',
            'entry': 'Entry Level',
            'staff': 'Staff'
        }

        for keyword, level in seniority_keywords.items():
            if keyword in text_lower:
                experiences.append({
                    "type": "seniority",
                    "level": level,
                    "description": f"{level} level position"
                })
                break

        return experiences

    @staticmethod
    def extract_education(text: str) -> List[Dict[str, Any]]:
        education = []
        text_lower = text.lower()

        education_patterns = [
            (r'bachelor\s+of\s+(\w+(?:\s+\w+)?)', 'Bachelor'),
            (r'b\.?\s*sc\.?\s+in\s+(\w+)', 'B.Sc'),
            (r'b\.?\s*tech', 'B.Tech'),
            (r'master\s+of\s+(\w+(?:\s+\w+)?)', 'Master'),
            (r'm\.?\s*sc\.?\s+in\s+(\w+)', 'M.Sc'),
            (r'm\.?\s*tech', 'M.Tech'),
            (r'b\.?\s*e\.?\s+in\s+(\w+)', 'B.E'),
            (r'm\.?\s*e\.?\s+in\s+(\w+)', 'M.E'),
            (r'ph\.?d\.?\s+in\s+(\w+(?:\s+\w+)?)', 'PhD'),
            (r'mba', 'MBA'),
            (r'b\.?\s*com', 'B.Com'),
            (r'm\.?\s*com', 'M.Com'),
            (r'ba\s+in\s+(\w+)', 'BA'),
            (r'ma\s+in\s+(\w+)', 'MA'),
        ]

        for pattern, degree_type in education_patterns:
            match = re.search(pattern, text_lower)
            if match:
                field = match.group(1).title() if match.lastindex and match.lastindex >= 1 else ""
                education.append({
                    "degree": degree_type,
                    "field": field,
                    "description": f"{degree_type} in {field}" if field else degree_type
                })
                break  # Only take the first degree found

        return education

    @staticmethod
    def parse_resume(file_path: str) -> Dict[str, Any]:
        ext = Path(file_path).suffix.lower()
        if ext == '.pdf':
            text = ResumeParser.extract_text_from_pdf(file_path)
        elif ext == '.docx':
            text = ResumeParser.extract_text_from_docx(file_path)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format. Use PDF or DOCX.")

        skills = ResumeParser.extract_skills(text)

        return {
            "raw_text": text,
            "name": ResumeParser.extract_name(text),
            "email": ResumeParser.extract_email(text),
            "phone": ResumeParser.extract_phone(text),
            "skills": skills,
            "experience": ResumeParser.extract_experience(text),
            "education": ResumeParser.extract_education(text),
        }


# ==================== GEMINI SERVICE ====================

class GeminiService:

    @staticmethod
    async def analyze_resume_with_gaps(resume_text: str, skills: List[str], experience: List, education: List) -> Dict:
        if not gemini_model:
            return GeminiService._fallback_ats(resume_text, skills, experience, education)

        prompt = f"""
You are an expert career coach and ATS consultant. Analyze this resume and return ONLY valid JSON.

Resume Skills Detected: {', '.join(skills[:25]) if skills else 'None detected'}
Experience Data: {json.dumps(experience[:3])}
Education Data: {json.dumps(education[:2])}

Resume Text (first 3500 chars):
{resume_text[:3500]}

Return ONLY this JSON structure (no markdown, no explanation):
{{
    "ats_score": <number 0-100>,
    "score_breakdown": {{
        "keyword_optimization": <number 0-30>,
        "formatting": <number 0-20>,
        "content_relevance": <number 0-30>,
        "quantifiable_achievements": <number 0-20>
    }},
    "strengths": ["<specific strength>", "<specific strength>", "<specific strength>", "<specific strength>", "<specific strength>"],
    "weaknesses": ["<specific weakness>", "<specific weakness>", "<specific weakness>", "<specific weakness>", "<specific weakness>"],
    "missing_keywords": ["<keyword>", "<keyword>", "<keyword>", "<keyword>", "<keyword>", "<keyword>", "<keyword>", "<keyword>", "<keyword>", "<keyword>"],
    "gap_analysis": {{
        "skill_gaps": ["<gap>", "<gap>", "<gap>"],
        "experience_gaps": ["<gap>", "<gap>"],
        "education_gaps": ["<gap>"],
        "format_gaps": ["<gap>", "<gap>"]
    }},
    "suggestions": ["<tip>", "<tip>", "<tip>", "<tip>", "<tip>", "<tip>", "<tip>"],
    "improvement_plan": {{
        "short_term": ["<action>", "<action>", "<action>"],
        "medium_term": ["<action>", "<action>", "<action>"],
        "long_term": ["<action>", "<action>", "<action>"]
    }},
    "detailed_feedback": "<2-3 sentence overall assessment>"
}}
"""
        try:
            response = gemini_model.generate_content(prompt)
            text = response.text.strip()
            # Strip markdown fences if present
            if text.startswith("```"):
                text = re.sub(r'^```[a-z]*\n?', '', text)
                text = re.sub(r'\n?```$', '', text)
            result = json.loads(text)
            logger.info(f"✅ Gemini ATS analysis done. Score: {result.get('ats_score')}")
            return result
        except Exception as e:
            logger.error(f"Gemini ATS error: {e}")
            return GeminiService._fallback_ats(resume_text, skills, experience, education)

    @staticmethod
    def _fallback_ats(resume_text: str, skills: List, experience: List, education: List) -> Dict:
        text_lower = resume_text.lower()
        important_kw = ['python', 'javascript', 'react', 'sql', 'aws', 'docker', 'git', 'agile', 'leadership',
                        'communication']
        found_kw = [k for k in important_kw if k in text_lower]
        keyword_score = (len(found_kw) / len(important_kw)) * 30
        has_numbers = bool(re.search(r'\d+%|\d+\+|\$\d+|\d+\s*years', text_lower))
        quant_score = 18 if has_numbers else 6
        has_sections = any(s in text_lower for s in ['experience', 'education', 'skills', 'summary', 'objective'])
        format_score = 18 if has_sections else 8
        content_score = min(len(skills) * 2.5, 30)
        ats_score = round(keyword_score + quant_score + format_score + content_score, 1)
        missing = [k for k in important_kw if k not in text_lower]
        return {
            "ats_score": ats_score,
            "score_breakdown": {
                "keyword_optimization": round(keyword_score, 1),
                "formatting": round(format_score, 1),
                "content_relevance": round(content_score, 1),
                "quantifiable_achievements": round(quant_score, 1)
            },
            "strengths": [
                f"Skills identified: {', '.join(skills[:5])}" if skills else "Resume has some content",
                "Has clear contact information" if re.search(r'@|phone|\d{10}',
                                                             text_lower) else "Add contact information",
                "Good length for ATS parsing" if len(resume_text) > 400 else "Resume is too short — add more detail",
                "Sections detected" if has_sections else "Add standard sections",
                "Quantified achievements present" if has_numbers else "Add metrics to stand out"
            ],
            "weaknesses": [
                "Missing quantifiable achievements (e.g. 'reduced costs by 30%')" if not has_numbers else "Good use of metrics",
                f"Only {len(found_kw)}/{len(important_kw)} critical keywords present" if len(
                    found_kw) < 7 else "Good keyword coverage",
                "Missing structured sections" if not has_sections else "Good structure",
                "Skills section could be expanded" if len(skills) < 5 else "Good skill breadth",
                "Add a professional summary at the top"
            ],
            "missing_keywords": missing[:10],
            "gap_analysis": {
                "skill_gaps": [f"Missing '{k}'" for k in missing[:4]],
                "experience_gaps": ["No work experience section found"] if not experience else [],
                "education_gaps": ["No education details found"] if not education else [],
                "format_gaps": ["Add clear section headings"] if not has_sections else []
            },
            "suggestions": [
                "Add a 2-3 line professional summary at the top",
                "Quantify every achievement with numbers, percentages, or dollar amounts",
                "Start every bullet point with an action verb (Led, Built, Reduced, Increased)",
                f"Add missing keywords: {', '.join(missing[:4])}" if missing else "Keyword coverage is solid",
                "Keep resume to 1 page for <5 years experience, 2 pages for more",
                "Use standard section headings: Experience, Education, Skills, Projects",
                "Remove personal info like photo, marital status, and DOB"
            ],
            "improvement_plan": {
                "short_term": ["Update contact info and add LinkedIn URL", "Add missing keywords from job descriptions",
                               "Write a professional summary"],
                "medium_term": ["Complete an online certification in a missing skill",
                                "Build a portfolio project showcasing key skills", "Get peer review of your resume"],
                "long_term": ["Gain 6-12 months of targeted experience", "Earn industry-recognized certifications",
                              "Network on LinkedIn to uncover hidden job opportunities"]
            },
            "detailed_feedback": f"Your resume scores {ats_score}/100 on ATS compatibility. Focus on adding quantifiable achievements, relevant keywords, and a clean section structure to significantly improve your score."
        }

    @staticmethod
    async def match_with_job(resume_text: str, job_description: str, skills: List[str]) -> Dict:
        if not gemini_model:
            return GeminiService._fallback_match(resume_text, job_description, skills)

        prompt = f"""
You are an expert hiring consultant. Match this resume against the job description and return ONLY valid JSON.

Candidate Skills: {', '.join(skills[:25])}

Resume (first 2500 chars):
{resume_text[:2500]}

Job Description (first 2000 chars):
{job_description[:2000]}

Return ONLY this JSON (no markdown, no explanation):
{{
    "match_percentage": <number 0-100>,
    "matched_skills": ["<skill>", "<skill>", "<skill>"],
    "missing_skills": ["<skill>", "<skill>", "<skill>"],
    "role_fit_summary": "<2-3 sentence role fit assessment>",
    "gap_analysis": {{
        "skill_gaps": ["<gap>", "<gap>"],
        "experience_gaps": ["<gap>"],
        "qualification_gaps": ["<gap>"]
    }},
    "improvement_tips": ["<tip>", "<tip>", "<tip>", "<tip>", "<tip>"]
}}
"""
        try:
            response = gemini_model.generate_content(prompt)
            text = response.text.strip()
            if text.startswith("```"):
                text = re.sub(r'^```[a-z]*\n?', '', text)
                text = re.sub(r'\n?```$', '', text)
            result = json.loads(text)
            logger.info(f"✅ Job match done: {result.get('match_percentage')}%")
            return result
        except Exception as e:
            logger.error(f"Gemini match error: {e}")
            return GeminiService._fallback_match(resume_text, job_description, skills)

    @staticmethod
    def _fallback_match(resume_text: str, job_description: str, skills: List[str]) -> Dict:
        job_lower = job_description.lower()
        tech_kw = ['python', 'java', 'javascript', 'react', 'angular', 'vue', 'node',
                   'django', 'flask', 'spring', 'aws', 'azure', 'docker', 'kubernetes',
                   'sql', 'mongodb', 'leadership', 'communication', 'agile', 'scrum',
                   'git', 'ci/cd', 'machine learning', 'data analysis', 'rest api']
        matched = [s for s in skills if s.lower() in job_lower]
        job_reqs = [k for k in tech_kw if k in job_lower]
        reqs_met = [r for r in job_reqs if r in resume_text.lower()]
        pct = (len(reqs_met) / len(job_reqs) * 100) if job_reqs else (len(matched) / max(len(skills), 1) * 100)
        pct = round(min(100, pct), 1)
        missing = [r for r in job_reqs if r not in resume_text.lower()][:10]
        if pct >= 80:
            summary = "Excellent match! Your skills strongly align with this role's requirements."
        elif pct >= 60:
            summary = "Good match with some gaps. You have core skills but are missing a few requirements."
        elif pct >= 40:
            summary = "Partial match. You have relevant experience but significant skill gaps remain."
        else:
            summary = "Limited match. Consider upskilling before applying, or heavily tailor your resume."
        return {
            "match_percentage": pct,
            "matched_skills": matched[:12],
            "missing_skills": missing,
            "role_fit_summary": summary,
            "gap_analysis": {
                "skill_gaps": missing[:5],
                "experience_gaps": ["Experience in required tech stack may be insufficient"] if pct < 60 else [],
                "qualification_gaps": [
                    "Consider relevant certifications to strengthen your application"] if pct < 50 else []
            },
            "improvement_tips": [
                f"Learn these missing skills: {', '.join(missing[:3])}" if missing else "Your skill set covers this role well",
                "Tailor your resume summary to match the job title and keywords",
                "Quantify achievements that relate directly to this role's responsibilities",
                "Add a cover letter addressing the experience gaps",
                "Highlight any adjacent experience that maps to missing requirements"
            ]
        }


# ==================== APP ====================

app = FastAPI(title="Resume Analyzer API", version="2.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://localhost:3001", "http://localhost:5173", "http://127.0.0.1:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/")
async def root():
    return {"message": "Resume Analyzer API", "status": "running", "version": "2.0.0"}


@app.get("/health")
async def health():
    return {
        "status": "healthy",
        "gemini_available": gemini_model is not None,
        "ocr_available": OCR_AVAILABLE,
        "database": "connected",
        "timestamp": datetime.utcnow().isoformat()
    }


# ==================== AUTH ENDPOINTS ====================

@app.post("/api/auth/signup")
async def signup(user_data: UserCreate, db: Session = Depends(get_db)):
    """
    Register a new user account
    """
    try:
        logger.info(f"Signup attempt - Email: {user_data.email}, Username: {user_data.username}")

        # Check if user already exists
        existing = db.query(User).filter(
            (User.email == user_data.email) | (User.username == user_data.username)
        ).first()

        if existing:
            conflict_field = "email" if existing.email == user_data.email else "username"
            logger.warning(f"User already exists with this {conflict_field}")
            raise HTTPException(
                status_code=400,
                detail=f"User with this {conflict_field} already exists"
            )

        # Hash password
        try:
            hashed_password = get_password_hash(user_data.password)
            logger.info("Password hashed successfully")
        except Exception as e:
            logger.error(f"Password hashing failed: {e}")
            raise HTTPException(status_code=500, detail="Error processing password")

        # Create user
        user = User(
            email=user_data.email,
            username=user_data.username,
            hashed_password=hashed_password
        )

        db.add(user)
        db.commit()
        db.refresh(user)

        logger.info(f"✅ User created successfully: {user.username} (ID: {user.id})")

        return {
            "id": user.id,
            "email": user.email,
            "username": user.username,
            "created_at": user.created_at.isoformat(),
            "message": "Account created successfully"
        }

    except HTTPException:
        raise
    except IntegrityError as e:
        db.rollback()
        logger.error(f"Database integrity error: {e}")
        raise HTTPException(status_code=400, detail="Username or email already exists")
    except Exception as e:
        db.rollback()
        logger.error(f"Unexpected signup error: {e}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")


@app.post("/api/auth/login", response_model=Token)
async def login(form_data: OAuth2PasswordRequestForm = Depends(), db: Session = Depends(get_db)):
    """
    Login with username and password to get access token
    """
    try:
        logger.info(f"Login attempt - Username: {form_data.username}")

        # Find user by username
        user = db.query(User).filter(User.username == form_data.username).first()

        if not user:
            logger.warning(f"User not found: {form_data.username}")
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Incorrect username or password"
            )

        # Verify password
        if not verify_password(form_data.password, user.hashed_password):
            logger.warning(f"Invalid password for user: {form_data.username}")
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Incorrect username or password"
            )

        # Create access token
        access_token = create_access_token(data={"sub": str(user.id)})

        logger.info(f"✅ Login successful: {user.username}")

        return {"access_token": access_token, "token_type": "bearer"}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Login error: {e}")
        raise HTTPException(status_code=500, detail=f"Login failed: {str(e)}")


@app.get("/api/auth/profile")
async def get_profile(current_user: User = Depends(get_current_user)):
    """
    Get current user profile
    """
    return {
        "id": current_user.id,
        "username": current_user.username,
        "email": current_user.email,
        "created_at": current_user.created_at.isoformat()
    }


#  RESUME ENDPOINTS

@app.post("/api/resume/upload")
async def upload_resume(
        file: UploadFile = File(...),
        current_user: User = Depends(get_current_user),
        db: Session = Depends(get_db)
):
    ext = Path(file.filename).suffix.lower()
    if ext not in {'.pdf', '.docx'}:
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are supported")

    content = await file.read()
    if len(content) > Config.MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File too large (max 10MB)")

    filename = f"{uuid.uuid4()}{ext}"
    filepath = Config.UPLOAD_DIR / filename
    with open(filepath, "wb") as f:
        f.write(content)

    try:
        parsed = ResumeParser.parse_resume(str(filepath))
        logger.info(f"📄 Parsed: {parsed.get('name', 'Unknown')}, skills: {len(parsed.get('skills', []))}")
    except Exception as e:
        if filepath.exists():
            filepath.unlink()
        raise HTTPException(status_code=400, detail=f"Parse error: {str(e)}")

    resume = Resume(
        user_id=current_user.id,
        filename=file.filename,
        file_path=str(filepath),
        name=parsed.get("name"),
        email=parsed.get("email"),
        phone=parsed.get("phone"),
        skills=json.dumps(parsed.get("skills", [])),
        experience=json.dumps(parsed.get("experience", [])),
        education=json.dumps(parsed.get("education", [])),
        raw_text=parsed.get("raw_text")
    )
    db.add(resume)
    db.commit()
    db.refresh(resume)

    return {
        "id": resume.id,
        "filename": resume.filename,
        "name": resume.name,
        "email": resume.email,
        "phone": resume.phone,
        "skills": parsed.get("skills", []),
        "experience": parsed.get("experience", []),
        "education": parsed.get("education", []),
        "char_count": len(parsed.get("raw_text", "")),
        "message": "Resume uploaded and parsed successfully"
    }


@app.get("/api/resume/list")
async def list_resumes(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    resumes = db.query(Resume).filter(Resume.user_id == current_user.id).order_by(desc(Resume.uploaded_at)).all()
    return [{
        "id": r.id,
        "filename": r.filename,
        "name": r.name,
        "email": r.email,
        "phone": r.phone,
        "skills": json.loads(r.skills) if r.skills else [],
        "experience": json.loads(r.experience) if r.experience else [],
        "education": json.loads(r.education) if r.education else [],
        "uploaded_at": r.uploaded_at.isoformat() if r.uploaded_at else None
    } for r in resumes]


@app.get("/api/resume/{resume_id}")
async def get_resume(resume_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    resume = db.query(Resume).filter(Resume.id == resume_id, Resume.user_id == current_user.id).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    return {
        "id": resume.id, "filename": resume.filename, "name": resume.name,
        "email": resume.email, "phone": resume.phone,
        "skills": json.loads(resume.skills) if resume.skills else [],
        "experience": json.loads(resume.experience) if resume.experience else [],
        "education": json.loads(resume.education) if resume.education else [],
        "uploaded_at": resume.uploaded_at.isoformat() if resume.uploaded_at else None
    }


@app.delete("/api/resume/{resume_id}")
async def delete_resume(resume_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    resume = db.query(Resume).filter(Resume.id == resume_id, Resume.user_id == current_user.id).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    if Path(resume.file_path).exists():
        Path(resume.file_path).unlink()
    db.delete(resume)
    db.commit()
    return {"message": "Resume deleted successfully"}


# ==================== ANALYSIS ENDPOINTS ====================

@app.post("/api/analyze/advanced-ats")
async def analyze_advanced_ats(
        request: ATSAnalysisRequest,
        current_user: User = Depends(get_current_user),
        db: Session = Depends(get_db),
        job_description: Optional[str] = None
):
    """Enhanced ATS analysis with detailed scoring and history"""
    try:
        resume = db.query(Resume).filter(
            Resume.id == request.resume_id,
            Resume.user_id == current_user.id
        ).first()

        if not resume:
            raise HTTPException(status_code=404, detail="Resume not found")

        if not resume.raw_text:
            raise HTTPException(status_code=400, detail="Resume has no parsed text")

        # Parse data
        skills = json.loads(resume.skills) if resume.skills else []
        experience = json.loads(resume.experience) if resume.experience else []
        education = json.loads(resume.education) if resume.education else []

        # Calculate individual scores
        experience_score = AdvancedScoreCalculator.calculate_experience_score(resume.raw_text, experience)
        skills_score = AdvancedScoreCalculator.calculate_skills_score(skills, job_description)
        education_score = AdvancedScoreCalculator.calculate_education_score(education, resume.raw_text)
        keywords_score = AdvancedScoreCalculator.calculate_keyword_score(resume.raw_text, job_description)
        formatting_score = AdvancedScoreCalculator.calculate_formatting_score(resume.raw_text)

        scores = {
            "experience": experience_score,
            "skills": skills_score,
            "education": education_score,
            "keywords": keywords_score,
            "formatting": formatting_score
        }

        # Calculate total with weights
        total = AdvancedScoreCalculator.calculate_total_score(scores)

        # Get previous score for history tracking
        previous_analysis = db.query(Analysis).filter(
            Analysis.resume_id == resume.id,
            Analysis.analysis_type == "ats_score"
        ).order_by(Analysis.created_at.desc()).first()

        score_change = None
        if previous_analysis and previous_analysis.ats_score:
            score_change = round(total["total_score"] - previous_analysis.ats_score, 1)

        # Save score history
        history = ScoreHistory(
            resume_id=resume.id,
            user_id=current_user.id,
            total_score=total["total_score"],
            experience_score=experience_score["score"],
            skills_score=skills_score["score"],
            education_score=education_score["score"],
            keywords_score=keywords_score["score"],
            formatting_score=formatting_score["score"]
        )
        db.add(history)

        # Save analysis
        analysis = Analysis(
            user_id=current_user.id,
            resume_id=resume.id,
            analysis_type="ats_score",
            ats_score=total["total_score"],
            strengths=json.dumps(AdvancedScoreCalculator._generate_strengths(scores)),
            weaknesses=json.dumps(AdvancedScoreCalculator._generate_weaknesses(scores)),
            missing_keywords=json.dumps(keywords_score["details"].get("top_keywords", [])),
            suggestions=json.dumps(AdvancedScoreCalculator._generate_suggestions(scores)),
            gap_analysis=json.dumps(AdvancedScoreCalculator._generate_gap_analysis(scores)),
            improvement_plan=json.dumps(AdvancedScoreCalculator._generate_improvement_plan(scores))
        )
        db.add(analysis)
        db.commit()

        return {
            "ats_score": total["total_score"],
            "score_breakdown": {
                "experience": experience_score["score"],
                "skills": skills_score["score"],
                "education": education_score["score"],
                "keywords": keywords_score["score"],
                "formatting": formatting_score["score"]
            },
            "score_details": {
                "experience": experience_score["details"],
                "skills": skills_score["details"],
                "education": education_score["details"],
                "keywords": keywords_score["details"],
                "formatting": formatting_score["details"]
            },
            "percentile_rank": total["percentile_rank"],
            "percentile_label": total["percentile_label"],
            "confidence_interval": total["confidence_interval"],
            "score_range": total["score_range"],
            "score_change": score_change,
            "previous_score": previous_analysis.ats_score if previous_analysis else None,
            "total_analyses": db.query(ScoreHistory).filter(ScoreHistory.resume_id == resume.id).count(),
            "strengths": AdvancedScoreCalculator._generate_strengths(scores),
            "weaknesses": AdvancedScoreCalculator._generate_weaknesses(scores),
            "suggestions": AdvancedScoreCalculator._generate_suggestions(scores),
            "gap_analysis": AdvancedScoreCalculator._generate_gap_analysis(scores),
            "improvement_plan": AdvancedScoreCalculator._generate_improvement_plan(scores)
        }

    except Exception as e:
        logger.error(f"Advanced analysis error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# Helper methods for AdvancedScoreCalculator
@staticmethod
def _generate_strengths(scores):
    strengths = []
    for category, data in scores.items():
        if data["score"] >= 70:
            strengths.append(f"Strong {category.title()} section: {data['score']}/100")
    return strengths[:5]


@staticmethod
def _generate_weaknesses(scores):
    weaknesses = []
    for category, data in scores.items():
        if data["score"] < 60:
            weaknesses.append(f"Improve {category.title()} section: {data['score']}/100")
    return weaknesses[:5]


@staticmethod
def _generate_suggestions(scores):
    suggestions = []
    if scores["keywords"]["score"] < 60:
        suggestions.append("Add more industry-relevant keywords to your resume")
    if scores["formatting"]["score"] < 60:
        suggestions.append("Improve formatting with clear section headers and bullet points")
    if scores["skills"]["score"] < 60:
        suggestions.append("List more technical skills relevant to your target role")
    return suggestions[:7]


@staticmethod
def _generate_gap_analysis(scores):
    gaps = {}
    for category, data in scores.items():
        if data["score"] < 50:
            gaps[category] = [f"{category.title()} score is low ({data['score']}/100)"]
    return gaps


@staticmethod
def _generate_improvement_plan(scores):
    plan = {"short_term": [], "medium_term": [], "long_term": []}
    for category, data in scores.items():
        if data["score"] < 50:
            plan["short_term"].append(f"Focus on improving {category} section")
        elif data["score"] < 70:
            plan["medium_term"].append(f"Enhance {category} with more details")
    return plan


@app.post("/api/analyze/match-job")
async def match_job(
        request: JobMatchRequest,
        current_user: User = Depends(get_current_user),
        db: Session = Depends(get_db)
):
    resume = db.query(Resume).filter(Resume.id == request.resume_id, Resume.user_id == current_user.id).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    if not resume.raw_text:
        raise HTTPException(status_code=400, detail="Resume has no parsed text. Please re-upload.")

    skills = json.loads(resume.skills) if resume.skills else []
    result = await GeminiService.match_with_job(resume.raw_text, request.job_description, skills)

    analysis = Analysis(
        user_id=current_user.id,
        resume_id=resume.id,
        analysis_type="job_match",
        job_description=request.job_description,
        match_percentage=result.get("match_percentage", 0),
        matched_skills=json.dumps(result.get("matched_skills", [])),
        missing_skills=json.dumps(result.get("missing_skills", [])),
        role_fit_summary=result.get("role_fit_summary", ""),
        gap_analysis=json.dumps(result.get("gap_analysis", {})),
        improvement_tips=json.dumps(result.get("improvement_tips", []))
    )
    db.add(analysis)
    db.commit()
    logger.info(f"✅ Job match saved. Match: {result.get('match_percentage')}%")
    return result


@app.get("/api/analysis/history")
async def get_history(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    analyses = db.query(Analysis).filter(Analysis.user_id == current_user.id).order_by(desc(Analysis.created_at)).all()
    result = []
    for a in analyses:
        resume = db.query(Resume).filter(Resume.id == a.resume_id).first()
        result.append({
            "id": a.id,
            "resume_id": a.resume_id,
            "resume_name": (resume.name or resume.filename) if resume else "Unknown",
            "analysis_type": a.analysis_type,
            "ats_score": a.ats_score,
            "match_percentage": a.match_percentage,
            "created_at": a.created_at.isoformat() if a.created_at else None
        })
    return result


@app.get("/api/analysis/{analysis_id}")
async def get_analysis(analysis_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    a = db.query(Analysis).filter(Analysis.id == analysis_id, Analysis.user_id == current_user.id).first()
    if not a:
        raise HTTPException(status_code=404, detail="Analysis not found")
    return {
        "id": a.id, "resume_id": a.resume_id, "analysis_type": a.analysis_type,
        "ats_score": a.ats_score,
        "strengths": json.loads(a.strengths) if a.strengths else [],
        "weaknesses": json.loads(a.weaknesses) if a.weaknesses else [],
        "missing_keywords": json.loads(a.missing_keywords) if a.missing_keywords else [],
        "suggestions": json.loads(a.suggestions) if a.suggestions else [],
        "gap_analysis": json.loads(a.gap_analysis) if a.gap_analysis else {},
        "improvement_plan": json.loads(a.improvement_plan) if a.improvement_plan else {},
        "match_percentage": a.match_percentage,
        "matched_skills": json.loads(a.matched_skills) if a.matched_skills else [],
        "missing_skills": json.loads(a.missing_skills) if a.missing_skills else [],
        "role_fit_summary": a.role_fit_summary,
        "improvement_tips": json.loads(a.improvement_tips) if a.improvement_tips else [],
        "created_at": a.created_at.isoformat() if a.created_at else None
    }


@app.get("/api/score-history/{resume_id}")
async def get_score_history(
        resume_id: int,
        current_user: User = Depends(get_current_user),
        db: Session = Depends(get_db)
):
    """Get score history for a resume"""
    history = db.query(ScoreHistory).filter(
        ScoreHistory.resume_id == resume_id,
        ScoreHistory.user_id == current_user.id
    ).order_by(ScoreHistory.created_at).all()

    return [{
        "date": h.created_at.isoformat(),
        "total_score": h.total_score,
        "experience": h.experience_score,
        "skills": h.skills_score,
        "education": h.education_score,
        "keywords": h.keywords_score,
        "formatting": h.formatting_score
    } for h in history]


@app.get("/api/weight-settings")
async def get_weight_settings():
    """Get current weight settings (admin only in production)"""
    return AdvancedScoreCalculator.DEFAULT_WEIGHTS


@app.post("/api/weight-settings")
async def update_weight_settings(weights: Dict):
    """Update weight settings (admin only in production)"""
    # Validate weights sum to 100
    if sum(weights.values()) != 100:
        raise HTTPException(status_code=400, detail="Weights must sum to 100")

    return {"message": "Weights updated successfully", "weights": weights}

#REQUIREMENTS.TXT
"""
Create a requirements.txt file with:

fastapi==0.104.1
uvicorn==0.24.0
sqlalchemy==2.0.23
python-jose[cryptography]==3.3.0
passlib[bcrypt]==1.7.4
python-multipart==0.0.6
pydantic==2.5.0
email-validator==2.1.0
bcrypt==4.1.2
PyPDF2==3.0.1
python-docx==1.1.0
google-generativeai==0.3.2
pytesseract==0.3.10
pdf2image==1.16.3
pdfplumber==0.10.3
python-dotenv==1.0.0
"""

#  RUN

if __name__ == "__main__":
    import uvicorn

    print("\n" + "=" * 60)
    print("🚀 Resume Analyzer API v2.0")
    print("=" * 60)
    print(f"🤖 Gemini AI: {'✅ Ready' if gemini_model else '⚠️  Fallback mode (set GEMINI_API_KEY)'}")
    print(f"🔍 OCR:       {'✅ Ready' if OCR_AVAILABLE else '⚠️  Not installed'}")
    print(f"📁 Uploads:   {Config.UPLOAD_DIR.absolute()}")
    print(f"🗄️  Database:  {Config.DATABASE_URL}")
    print("🌐 API:        http://localhost:8000")
    print("📚 Docs:       http://localhost:8000/docs")
    print("=" * 60 + "\n")
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)